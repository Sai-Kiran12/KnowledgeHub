import logging
from pathlib import Path

import pandas as pd
import textract
from docx import Document as DocxDocument
from pypdf import PdfReader

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    '.txt',
    '.md',
    '.pdf',
    '.doc',
    '.docx',
    '.csv',
    '.png',
    '.jpg',
    '.jpeg',
    '.webp',
}

IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.webp'}


def load_document(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    logger.info('Loading document path=%s ext=%s', file_path, ext)

    if ext not in SUPPORTED_EXTENSIONS:
        logger.warning('Unsupported document extension path=%s ext=%s', file_path, ext)
        raise ValueError(f'Unsupported file extension: {ext}')

    if ext in {'.txt', '.md'}:
        text = file_path.read_text(encoding='utf-8', errors='ignore')
        logger.info('Loaded text/markdown chars=%s', len(text))
        return text

    if ext == '.pdf':
        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or '' for page in reader.pages]
        text = '\n'.join(pages)
        logger.info('Loaded pdf pages=%s chars=%s', len(reader.pages), len(text))
        return text

    if ext == '.docx':
        doc = DocxDocument(str(file_path))
        text = '\n'.join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
        logger.info('Loaded docx paragraphs=%s chars=%s', len(doc.paragraphs), len(text))
        return text

    if ext == '.doc':
        raw = textract.process(str(file_path))
        text = raw.decode('utf-8', errors='ignore')
        logger.info('Loaded doc chars=%s', len(text))
        return text

    if ext == '.csv':
        df = pd.read_csv(file_path)
        text = df.to_csv(index=False)
        logger.info('Loaded csv rows=%s cols=%s chars=%s', len(df.index), len(df.columns), len(text))
        return text

    return ''
